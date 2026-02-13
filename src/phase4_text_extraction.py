"""
Phase 4: Text Extraction & Structuring

This module handles:
- PDF text extraction
- OCR for scanned documents (images)
- Text normalization (fixing encoding issues like Ã³ → ó)
- Data extraction (names, dates, RUT, registry numbers)
- Structured JSON output

This prepares extracted data for Phase 5 (validation).
"""

from dataclasses import dataclass, field
from typing import List, Dict, Optional, Any
from pathlib import Path
from datetime import datetime
import json
import re

from src.phase3_document_intake import (
    DocumentCollection,
    DocumentType,
    DocumentTypeDetector,
    FileFormat,
    UploadedDocument,
)


class TextNormalizer:
    """
    Normalizes text extracted from documents.
    Fixes encoding issues common in OCR (especially Spanish tildes).
    """

    # Common OCR encoding errors in Spanish
    ENCODING_FIXES = {
        'Ã³': 'ó',
        'Ã¡': 'á',
        'Ã©': 'é',
        'Ã­': 'í',
        'Ãº': 'ú',
        'Ã±': 'ñ',
        'Ã': 'Á',
        'Ã‰': 'É',
        'Ã': 'Í',
        'Ã"': 'Ó',
        'Ãš': 'Ú',
        'Ñ': 'Ñ',  # Capital N with tilde
        'Â°': '°',
        'Âº': 'º',
        'Âª': 'ª',
    }

    @staticmethod
    def fix_encoding(text: str) -> str:
        """Fix common encoding issues in OCR text"""
        if not text:
            return text

        fixed_text = text
        for wrong, correct in TextNormalizer.ENCODING_FIXES.items():
            fixed_text = fixed_text.replace(wrong, correct)

        return fixed_text

    @staticmethod
    def strip_bom(text: str) -> str:
        """Strip Unicode BOM or visible UTF-8 BOM artifacts."""
        if not text:
            return text
        if text.startswith("\ufeff"):
            text = text.lstrip("\ufeff")
        # Sometimes BOM bytes show up as visible characters when mis-decoded.
        if text.startswith("ï»¿"):
            text = text[len("ï»¿") :]
        return text

    @staticmethod
    def _mojibake_score(text: str) -> int:
        """Heuristic score: higher means more mojibake-like artifacts."""
        if not text:
            return 0
        patterns = ("Ã", "Â", "ï»¿", "�")
        return sum(text.count(p) for p in patterns)

    @staticmethod
    def repair_mojibake(text: str) -> str:
        """
        Best-effort repair for common mojibake where UTF-8 bytes were decoded as latin-1/cp1252.
        Applies only when it improves the text by a simple heuristic.
        """
        if not text:
            return text

        original = text
        original_score = TextNormalizer._mojibake_score(original)
        if original_score == 0:
            return original

        candidates = []
        for codec in ("latin1", "cp1252"):
            try:
                repaired = original.encode(codec, errors="strict").decode("utf-8", errors="strict")
            except Exception:
                continue
            candidates.append(repaired)

        best = original
        best_score = original_score
        for cand in candidates:
            score = TextNormalizer._mojibake_score(cand)
            if score < best_score:
                best = cand
                best_score = score

        return best

    @staticmethod
    def normalize_whitespace(text: str) -> str:
        """Normalize whitespace (multiple spaces, tabs, newlines)"""
        if not text:
            return text

        # Replace multiple spaces with single space
        text = re.sub(r'\s+', ' ', text)
        # Remove leading/trailing whitespace
        text = text.strip()

        return text

    @staticmethod
    def normalize_text(text: str) -> str:
        """Apply all normalization steps"""
        text = TextNormalizer.strip_bom(text)
        text = TextNormalizer.repair_mojibake(text)
        text = TextNormalizer.fix_encoding(text)
        text = TextNormalizer.normalize_whitespace(text)
        return text


class DataExtractor:
    """
    Extracts structured data from text.
    Uses regex patterns to find names, dates, RUT numbers, etc.
    """

    # Regex patterns for Uruguayan documents
    PATTERNS = {
        'rut': r'\b\d{12}\b|\b\d{2}[\s\.-]?\d{3}[\s\.-]?\d{3}[\s\.-]?\d{4}\b',
        'ci': r'\b\d\.\d{3}\.\d{3}[-\s]?\d\b',  # Uruguayan CI format
        'date': r'\b\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b',
        'phone': r'\b\d{4}[-\s]?\d{4}\b|\b\d{3}[-\s]?\d{3}[-\s]?\d{3}\b',
        'email': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
        'registro_comercio': r'Registro\s+(?:de\s+)?Comercio\s+(?:N°|Nro\.?|Número)?\s*(\d+)',
        'acta_number': r'Acta\s+(?:N°|Nro\.?|Número)?\s*(\d+)',
        'padron_bps': r'Padrón\s+(?:BPS\s+)?(?:N°|Nro\.?|Número)?\s*(\d+)',
    }

    @staticmethod
    def extract_rut(text: str) -> Optional[str]:
        """Extract RUT (tax ID) from text"""
        match = re.search(DataExtractor.PATTERNS['rut'], text, re.IGNORECASE)
        if match:
            # Normalize RUT format (remove spaces, dots, dashes)
            rut = re.sub(r'[\s\.-]', '', match.group(0))
            return rut
        return None

    @staticmethod
    def extract_ci(text: str) -> Optional[str]:
        """Extract CI (Cédula de Identidad) from text"""
        match = re.search(DataExtractor.PATTERNS['ci'], text, re.IGNORECASE)
        return match.group(0) if match else None

    @staticmethod
    def extract_dates(text: str) -> List[str]:
        """Extract all dates from text"""
        matches = re.findall(DataExtractor.PATTERNS['date'], text)
        return matches

    @staticmethod
    def extract_emails(text: str) -> List[str]:
        """Extract email addresses from text"""
        matches = re.findall(DataExtractor.PATTERNS['email'], text)
        return matches

    @staticmethod
    def extract_registro_comercio(text: str) -> Optional[str]:
        """Extract Registro de Comercio number"""
        match = re.search(DataExtractor.PATTERNS['registro_comercio'], text, re.IGNORECASE)
        if match:
            return match.group(1) if match.lastindex >= 1 else match.group(0)
        return None

    @staticmethod
    def extract_acta_number(text: str) -> Optional[str]:
        """Extract Acta number"""
        match = re.search(DataExtractor.PATTERNS['acta_number'], text, re.IGNORECASE)
        if match:
            return match.group(1) if match.lastindex >= 1 else match.group(0)
        return None

    @staticmethod
    def extract_padron_bps(text: str) -> Optional[str]:
        """Extract Padrón BPS number"""
        match = re.search(DataExtractor.PATTERNS['padron_bps'], text, re.IGNORECASE)
        if match:
            return match.group(1) if match.lastindex >= 1 else match.group(0)
        return None

    @staticmethod
    def extract_company_name(text: str) -> Optional[str]:
        """
        Extract company name (S.A., S.R.L., etc.)
        This is a simple heuristic - can be improved with NER
        """
        # Look for patterns like "NOMBRE S.A." or "NOMBRE SOCIEDAD ANÓNIMA"
        patterns = [
            r'([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑa-záéíóúñ\s]+)\s+S\.?A\.?(?:\s|$)',
            r'([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑa-záéíóúñ\s]+)\s+SOCIEDAD\s+ANÓNIMA',
            r'([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑa-záéíóúñ\s]+)\s+S\.?R\.?L\.?(?:\s|$)',
        ]

        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                company_name = match.group(0).strip()
                return company_name

        return None


@dataclass
class ExtractedData:
    """Structured data extracted from a document"""
    document_type: DocumentType
    raw_text: str
    normalized_text: str

    # Extracted fields
    company_name: Optional[str] = None
    rut: Optional[str] = None
    ci: Optional[str] = None
    registro_comercio: Optional[str] = None
    acta_number: Optional[str] = None
    padron_bps: Optional[str] = None
    dates: List[str] = field(default_factory=list)
    emails: List[str] = field(default_factory=list)

    # Specific date fields for validation
    company_constitution_date: Optional[str] = None
    statute_approval_date: Optional[str] = None
    registration_date: Optional[str] = None
    acta_date: Optional[str] = None
    registry_number: Optional[str] = None

    # Metadata
    extraction_timestamp: datetime = field(default_factory=datetime.now)
    extraction_method: str = "text"  # "text" or "ocr"
    confidence: float = 1.0  # 0.0 to 1.0

    # Additional structured data
    additional_fields: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> dict:
        return {
            "document_type": self.document_type.value if self.document_type else None,
            "company_name": self.company_name,
            "rut": self.rut,
            "ci": self.ci,
            "registro_comercio": self.registro_comercio,
            "acta_number": self.acta_number,
            "padron_bps": self.padron_bps,
            "dates": self.dates,
            "emails": self.emails,
            "company_constitution_date": self.company_constitution_date,
            "statute_approval_date": self.statute_approval_date,
            "registration_date": self.registration_date,
            "acta_date": self.acta_date,
            "registry_number": self.registry_number,
            "extraction_timestamp": self.extraction_timestamp.isoformat(),
            "extraction_method": self.extraction_method,
            "confidence": self.confidence,
            "additional_fields": self.additional_fields,
            "text_preview": self.normalized_text[:200] + "..." if len(self.normalized_text) > 200 else self.normalized_text
        }

    def to_json(self) -> str:
        return json.dumps(self.to_dict(), ensure_ascii=False, indent=2)

    def get_summary(self) -> str:
        """Get human-readable summary"""
        doc_type = self.document_type.value if self.document_type else "desconocido"

        summary = f"""
📄 DOCUMENTO: {doc_type.upper()}
   Método extracción: {self.extraction_method.upper()}
   Confianza: {self.confidence * 100:.1f}%

📊 DATOS EXTRAÍDOS:
"""
        if self.company_name:
            summary += f"   Empresa: {self.company_name}\n"
        if self.rut:
            summary += f"   RUT: {self.rut}\n"
        if self.ci:
            summary += f"   CI: {self.ci}\n"
        if self.registro_comercio:
            summary += f"   Registro Comercio: {self.registro_comercio}\n"
        if self.acta_number:
            summary += f"   Acta N°: {self.acta_number}\n"
        if self.padron_bps:
            summary += f"   Padrón BPS: {self.padron_bps}\n"
        if self.dates:
            summary += f"   Fechas encontradas: {', '.join(self.dates[:3])}\n"
        if self.emails:
            summary += f"   Emails: {', '.join(self.emails)}\n"

        return summary


@dataclass
class DocumentExtractionResult:
    """Result of extracting data from a single document"""
    document: UploadedDocument
    extracted_data: Optional[ExtractedData] = None
    error: Optional[str] = None
    success: bool = False

    def to_dict(self) -> dict:
        return {
            "file_name": self.document.file_name,
            "file_path": str(self.document.file_path),
            "document_type": self.document.detected_type.value if self.document.detected_type else None,
            "success": self.success,
            "error": self.error,
            "extracted_data": self.extracted_data.to_dict() if self.extracted_data else None
        }


@dataclass
class CollectionExtractionResult:
    """Result of extracting data from entire collection"""
    collection: DocumentCollection
    extraction_results: List[DocumentExtractionResult] = field(default_factory=list)
    extraction_timestamp: datetime = field(default_factory=datetime.now)

    def get_success_count(self) -> int:
        """Get count of successful extractions"""
        return sum(1 for result in self.extraction_results if result.success)

    def get_failed_count(self) -> int:
        """Get count of failed extractions"""
        return sum(1 for result in self.extraction_results if not result.success)

    def to_dict(self) -> dict:
        return {
            "total_documents": len(self.extraction_results),
            "successful": self.get_success_count(),
            "failed": self.get_failed_count(),
            "extraction_timestamp": self.extraction_timestamp.isoformat(),
            "results": [result.to_dict() for result in self.extraction_results]
        }

    def to_json(self) -> str:
        return json.dumps(self.to_dict(), ensure_ascii=False, indent=2)

    def get_summary(self) -> str:
        """Get human-readable summary"""
        total = len(self.extraction_results)
        success = self.get_success_count()
        failed = self.get_failed_count()
        success_rate = (success / total * 100) if total > 0 else 0

        summary = f"""
╔══════════════════════════════════════════════════════════════╗
║              EXTRACCIÓN DE DATOS - FASE 4                    ║
╚══════════════════════════════════════════════════════════════╝

📊 RESUMEN:
   Total documentos: {total}
   Exitosos: {success}
   Fallidos: {failed}
   Tasa de éxito: {success_rate:.1f}%

📄 RESULTADOS POR DOCUMENTO:
"""
        for result in self.extraction_results:
            status = "✅" if result.success else "❌"
            summary += f"\n{status} {result.document.file_name}\n"
            if result.success and result.extracted_data:
                summary += result.extracted_data.get_summary()
            elif result.error:
                summary += f"   Error: {result.error}\n"

        return summary


class TextExtractor:
    """
    Main service class for text extraction.
    Handles both digital PDFs and scanned documents (OCR).
    """

    @staticmethod
    def extract_from_text_file(file_path: Path) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with latin-1 encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()

    @staticmethod
    def extract_from_pdf(file_path: Path) -> str:
        """
        Extract text from PDF file.

        Uses PyPDF2 for digital PDFs and falls back to OCR when needed.
        """
        try:
            from PyPDF2 import PdfReader
        except Exception as exc:
            raise RuntimeError("PyPDF2 is required for PDF extraction.") from exc

        reader = PdfReader(str(file_path))
        text_chunks = []
        for page in reader.pages:
            text_chunks.append(page.extract_text() or "")
        text = "\n".join(text_chunks).strip()

        if text:
            return text

        # Fallback to OCR if PDF text is empty
        return TextExtractor.extract_from_pdf_ocr(file_path)

    @staticmethod
    def extract_from_docx(file_path: Path) -> str:
        """
        Extract text from DOCX file.

        Uses python-docx for Word documents.
        """
        try:
            from docx import Document
        except Exception as exc:
            raise RuntimeError("python-docx is required for DOCX extraction.") from exc

        doc = Document(str(file_path))
        paragraphs = [para.text for para in doc.paragraphs if para.text]
        return "\n".join(paragraphs)

    @staticmethod
    def extract_from_image_ocr(file_path: Path) -> str:
        """
        Extract text from image using OCR.

        Uses pytesseract with Spanish language support when available.
        """
        try:
            from PIL import Image
            import pytesseract
        except Exception as exc:
            raise RuntimeError("pytesseract and Pillow are required for OCR.") from exc

        image = Image.open(file_path)
        try:
            return pytesseract.image_to_string(image, lang="spa")
        except Exception:
            return pytesseract.image_to_string(image)

    @staticmethod
    def extract_from_pdf_ocr(file_path: Path) -> str:
        """
        Extract text from a scanned PDF using OCR.
        Requires pdf2image and pytesseract.
        """
        try:
            from pdf2image import convert_from_path
            import pytesseract
        except Exception as exc:
            raise RuntimeError("pdf2image and pytesseract are required for PDF OCR.") from exc

        images = convert_from_path(str(file_path))
        text_chunks = []
        for image in images:
            try:
                text_chunks.append(pytesseract.image_to_string(image, lang="spa"))
            except Exception:
                text_chunks.append(pytesseract.image_to_string(image))
        return "\n".join(text_chunks).strip()

    @staticmethod
    def extract_text(document: UploadedDocument) -> tuple[str, str]:
        """
        Extract text from document based on file format.

        Returns:
            tuple: (raw_text, extraction_method)
        """
        file_path = document.file_path

        if document.file_format == FileFormat.TXT:
            text = TextExtractor.extract_from_text_file(file_path)
            return text, "text"

        elif document.file_format == FileFormat.PDF:
            if document.is_scanned:
                text = TextExtractor.extract_from_image_ocr(file_path)
                return text, "ocr"
            else:
                text = TextExtractor.extract_from_pdf(file_path)
                return text, "text"

        elif document.file_format in [FileFormat.DOCX, FileFormat.DOC]:
            text = TextExtractor.extract_from_docx(file_path)
            return text, "text"

        elif document.file_format in [FileFormat.JPG, FileFormat.JPEG, FileFormat.PNG]:
            text = TextExtractor.extract_from_image_ocr(file_path)
            return text, "ocr"

        else:
            raise ValueError(f"Unsupported file format: {document.file_format}")

    @staticmethod
    def process_document(document: UploadedDocument) -> DocumentExtractionResult:
        """
        Process a single document: extract text and structure data.

        Args:
            document: UploadedDocument to process

        Returns:
            DocumentExtractionResult
        """
        try:
            # Step 1: Extract raw text
            raw_text, extraction_method = TextExtractor.extract_text(document)

            # Step 2: Normalize text
            normalized_text = TextNormalizer.normalize_text(raw_text)

            # Step 2b: If filename-based detection failed, try content-based detection.
            if not document.detected_type and normalized_text:
                inferred_type = DocumentTypeDetector.detect_from_text(normalized_text)
                if inferred_type:
                    document.detected_type = inferred_type
                    if isinstance(document.metadata, dict):
                        document.metadata["detected_type_source"] = "content"

            # Step 3: Extract structured data
            extracted_data = ExtractedData(
                document_type=document.detected_type,
                raw_text=raw_text,
                normalized_text=normalized_text,
                extraction_method=extraction_method
            )

            # Extract specific fields
            extracted_data.company_name = DataExtractor.extract_company_name(normalized_text)
            extracted_data.rut = DataExtractor.extract_rut(normalized_text)
            extracted_data.ci = DataExtractor.extract_ci(normalized_text)
            extracted_data.registro_comercio = DataExtractor.extract_registro_comercio(normalized_text)
            extracted_data.acta_number = DataExtractor.extract_acta_number(normalized_text)
            extracted_data.padron_bps = DataExtractor.extract_padron_bps(normalized_text)
            extracted_data.dates = DataExtractor.extract_dates(normalized_text)
            extracted_data.emails = DataExtractor.extract_emails(normalized_text)

            # Set confidence (OCR = lower confidence)
            extracted_data.confidence = 0.8 if extraction_method == "ocr" else 1.0

            return DocumentExtractionResult(
                document=document,
                extracted_data=extracted_data,
                success=True
            )

        except Exception as e:
            return DocumentExtractionResult(
                document=document,
                error=str(e),
                success=False
            )

    @staticmethod
    def process_collection(collection: DocumentCollection) -> CollectionExtractionResult:
        """
        Process entire document collection.

        Args:
            collection: DocumentCollection to process

        Returns:
            CollectionExtractionResult
        """
        result = CollectionExtractionResult(collection=collection)

        for document in collection.documents:
            extraction_result = TextExtractor.process_document(document)
            result.extraction_results.append(extraction_result)

        return result

    @staticmethod
    def save_extraction_result(result: CollectionExtractionResult, output_path: str) -> None:
        """Save extraction result to JSON file"""
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(result.to_json())
        print(f"\n✅ Resultados de extracción guardados en: {output_path}")


def example_usage():
    """Example usage of Phase 4"""

    print("\n" + "="*70)
    print("  EJEMPLOS DE USO - FASE 4: EXTRACCIÓN DE TEXTO")
    print("="*70)

    # Example 1: Text normalization
    print("\n📌 Ejemplo 1: Normalización de texto con errores de encoding")
    print("-" * 70)

    raw_text = "ResoluciÃ³n del Directorio de la empresa GIRTEC S.A."
    normalized = TextNormalizer.normalize_text(raw_text)
    print(f"Original:    {raw_text}")
    print(f"Normalizado: {normalized}")

    # Example 2: Data extraction
    print("\n\n📌 Ejemplo 2: Extracción de datos estructurados")
    print("-" * 70)

    sample_text = """
    GIRTEC SOCIEDAD ANÓNIMA
    RUT: 21 234 567 8901
    Registro de Comercio Nro. 12345
    Acta N° 45 del 15/06/2023
    Padrón BPS: 98765
    Email: contacto@girtec.com.uy
    """

    print(f"Texto muestra:\n{sample_text}\n")
    print("Datos extraídos:")
    print(f"  Empresa: {DataExtractor.extract_company_name(sample_text)}")
    print(f"  RUT: {DataExtractor.extract_rut(sample_text)}")
    print(f"  Registro: {DataExtractor.extract_registro_comercio(sample_text)}")
    print(f"  Acta: {DataExtractor.extract_acta_number(sample_text)}")
    print(f"  Padrón BPS: {DataExtractor.extract_padron_bps(sample_text)}")
    print(f"  Fechas: {DataExtractor.extract_dates(sample_text)}")
    print(f"  Emails: {DataExtractor.extract_emails(sample_text)}")

    # Example 3: Processing a mock document
    print("\n\n📌 Ejemplo 3: Procesamiento de documento completo")
    print("-" * 70)

    # Create a mock document
    from src.phase3_document_intake import UploadedDocument, FileFormat, DocumentType
    from datetime import datetime

    mock_doc = UploadedDocument(
        file_path=Path("/mock/estatuto_girtec.txt"),
        file_name="estatuto_girtec.txt",
        file_format=FileFormat.TXT,
        file_size_bytes=1024,
        upload_timestamp=datetime.now(),
        detected_type=DocumentType.ESTATUTO,
        is_scanned=False
    )

    # Note: This would fail because file doesn't exist
    # In real usage, you'd have actual files
    print("⚠️  Para procesar documentos reales, use:")
    print("   result = TextExtractor.process_document(document)")
    print("   print(result.extracted_data.get_summary())")


if __name__ == "__main__":
    example_usage()
